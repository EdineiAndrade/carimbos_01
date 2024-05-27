from datetime import datetime
from dateutil.relativedelta import relativedelta
from num2words import num2words
import pandas as pd
from docx import Document
from docx.shared import Pt
import locale
import time
import os

# Definir a localidade para o Brasil
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

# Carregar dados do arquivo Excel
#pasta 
pasta= 'C:\\carimbos_res'
pasta_modelos = 'C:\\carimbos_res\\modelos'
pasta_carimbo = 'C:\\carimbos_res\\carimbos'

arquivo_excel = os.path.join(pasta,'lista_clientes.xlsx')
arquivo_usual = os.path.join(pasta_modelos,'carimbo_usual.docx')
arquivo_investimento = os.path.join(pasta_modelos,'carimbo_res_inv.docx')
arquivo_custeio = os.path.join(pasta_modelos,'carimbo_res_inv.docx')

df = pd.read_excel(arquivo_excel)

quantidade_linhas = df.shape[0]

#tempo inicial
tempo_inicio = time.time()
data_hora_atual1 = datetime.now()
data_format = data_hora_atual1.strftime("%Y/%m/%d")
data_format = f'{data_format[8:10]}/{data_format[5:7]}/{data_format[:4]}'
hora_format = data_hora_atual1.strftime("%H:%M:%S")
print(f'===========================================================================')    
print(f'=============| GERADOR DE CARIMBOS INICIADO | {data_format} {hora_format}|=====') 
print(f'===========================================================================\n')
print(f'=====[ Verificando informações...\n')

# Loop pelos clientes
for index, row in df.iterrows():
    time.sleep(1)
    # Extrair informações do DataFrame
    agencia = row['AGENCIA'] 
    agencia = agencia.title()
    uf = row['ESTADO']
    agencia_uf = f'{agencia}-{uf}'     
    cpf = int(row['CPF'])
    cpf = "{:011d}".format(cpf)
    cpf = f'{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}'
    cliente = row['CLIENTE']
    cliente = cliente.title()
    contrato = row['CONTRATO']
    ccb = row['CCB']
    n_parcelas = row['QTD PARCELAS']
    valor_extenso = num2words(n_parcelas, lang='pt_BR')
    if n_parcelas == 1:
         text_parcelas = f'1 (uma) parcela'
    elif n_parcelas == 2:
        text_parcelas = f'2 (duas) parcelas'      
    elif n_parcelas > 2:
        text_parcelas = f'{int(n_parcelas)} ({valor_extenso}) parcelas'     
    periodicidade = row['PERIODICIDADE'].upper()
    if periodicidade == 'ANUAL':
         text_periodicidade = 'anuais'
    elif periodicidade == 'SEMESTRAL':
        text_periodicidade = 'semestrais'
    tipo_reneg = row['TIPO_RENEG'].upper()
    if tipo_reneg == 'USUAL':
        modelo_word = Document(arquivo_usual)
    elif tipo_reneg == 'RESOLUCAO_I':
        modelo_word = Document(arquivo_investimento)
    elif tipo_reneg == 'RESOLUCAO_C':
        modelo_word = Document(arquivo_custeio)

    valor_renegociado = row['VALOR RENEGOCIADO']
    valor_renegociado_f = locale.format_string('%.2f', valor_renegociado, grouping=True)
    data_aditivo = row['DATA ADITIVO']
    if type(data_aditivo).__name__ == 'datetime' or type(data_aditivo).__name__ == 'Timestamp':
           data_aditivo = data_aditivo.strftime('%d/%m/%Y') 
    data_renegociacao = row['DATA RENEGOCIACAO']    
    if type(data_renegociacao).__name__ == 'datetime' or type(data_renegociacao).__name__ == 'Timestamp':
        data_renegociacao = data_renegociacao.strftime('%d/%m/%Y')     
    data_parcela1 = row['PARCELA 1']
    if type(data_parcela1).__name__ == 'datetime' or type(data_parcela1).__name__ == 'Timestamp':
        data_parcela1 = data_parcela1.strftime('%d/%m/%Y')
    local_data = f'{agencia_uf}, {data_aditivo}'         
    
    #criar a pasta da unidade
    if index == quantidade_linhas:
        print(f'Processo Finalizado....')
        break
    # Substituir palavras-chave no modelo do Word
    for i, paragraph in enumerate(modelo_word.paragraphs):
            #variaveis fixas(de acordo excel)
            if i == 0:
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#CLIENTE', cliente)
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#CPF', cpf)
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#CCB', ccb)
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#DATA_ADITIVO', data_aditivo)
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#VALOR_RENEGOCIADO', f'R$ {valor_renegociado_f}')    
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#N_PARCELAS', text_parcelas)
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#PERIODICIDADE', text_periodicidade)              
                modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace('#PARCELA1', data_parcela1) 
                modelo_word.paragraphs[40].text = modelo_word.paragraphs[40].text.replace('#LOCAL_DATA', local_data)        
            #valor das parcelas
            resto = (round(valor_renegociado,2)*100) % n_parcelas
            valor_demais_parcelas = (((valor_renegociado * 100) - resto) / n_parcelas)/100
            valor_demais_parcelas = locale.format_string('%.2f', valor_demais_parcelas, grouping=True)
            valor_demais_parcelas = f'{valor_demais_parcelas};'
            valor_ultima_parcela = float((((((valor_renegociado * 100) - resto) / n_parcelas ) + resto)/100))  
            valor_ultima_parcela = locale.format_string('%.2f', valor_ultima_parcela, grouping=True)
            valor_ultima_parcela = f'{valor_ultima_parcela}.'         
            reembolso_atual = f'Em {data_parcela1} R$ {valor_demais_parcelas}'
            if i == 0:
                modelo_word.paragraphs[5].text = modelo_word.paragraphs[5].text.replace('#REEMBOLSO1', reembolso_atual)
            if n_parcelas == 1 and '#REEMBOLSO' in modelo_word.paragraphs[6].text:
                parcela = 1
                for p in range(36,5,-1):
                        if p == 6: 
                             modelo_word.paragraphs[6]._element.clear()
                             break
                        paragraph_atual = modelo_word.paragraphs[p]
                        text_paragraph = paragraph_atual.text
                        if p < 12:                   
                            paragraph_atual._element.clear()
                        elif p >= 12 and '#REEMBOLSO' in text_paragraph:                             
                             modelo_word._element.body.remove(modelo_word.paragraphs[p]._element)
            elif n_parcelas > 1:
                    parcela = 1
                    for p in range(6,37):
                        parcela = parcela + 1                        
                        try:
                            paragraph_atual = modelo_word.paragraphs[p]
                            if periodicidade == 'SEMESTRAL':
                                if parcela == 2:
                                    data_atual = datetime.strptime(data_parcela1, "%d/%m/%Y")
                                    data_atual = data_atual +  relativedelta(months=6)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                    paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)

                                elif parcela >2:
                                    data_atual = data_atual +  relativedelta(months=6)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    if parcela == n_parcelas:
                                            reembolso_atual = f'Em {data_atual_f} R$ {valor_ultima_parcela}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                            continue
                                    elif parcela > n_parcelas and i ==0:
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
                                            if parcela > 11:
                                                if n_parcelas < 11:
                                                    remover = 18
                                                elif n_parcelas > 11:
                                                    remover = 18 - (n_parcelas-11)
                                                for r in range(1,remover):
                                                    if '#REEMBOLSO' in modelo_word.paragraphs[17].text:
                                                        modelo_word._element.body.remove(modelo_word.paragraphs[17]._element)                                           
                                                
                                    reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                    paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                    
                                      
                            elif periodicidade == 'ANUAL':                                        
                                        data_atual = f'{data_parcela1[:6]}{int(data_parcela1[6:]) + (parcela - 1)}'
                                        if parcela == n_parcelas:
                                            reembolso_atual = f'Em {data_atual} R$ {valor_ultima_parcela}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                            continue
                                        elif parcela > n_parcelas and i ==0:
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
                                            if parcela > 11:
                                                if n_parcelas < 11:
                                                    remover = 18
                                                elif n_parcelas > 11:
                                                    remover = 18 - (n_parcelas-11)
                                                for r in range(1,remover):
                                                    if '#REEMBOLSO' in modelo_word.paragraphs[17].text:
                                                        modelo_word._element.body.remove(modelo_word.paragraphs[17]._element)                     
                                        reembolso_atual = f'Em {data_atual} R$ {valor_demais_parcelas}'
                                        paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                        except:
                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
    # Modificar a fonte de todos os estilos no documento
    for style in modelo_word.styles:
        if style.type == 1:  # type 1 significa estilo de parágrafo
            font = style.font
            font.name = 'Times New Roman'  
            font.size = Pt(11)   # Definindo o tamanho da fonte para 11 pt
    # Salvar o documento personalizado com o nome do cliente
    modelo_word.save(f'{pasta_carimbo}\{cliente}.docx')
    
    restante = quantidade_linhas - index
    tempo_atual = time.time() - tempo_inicio 
    duracao = f"{int(tempo_atual // 3600):02d}:{int((tempo_atual % 3600) // 60):02d}:{int(tempo_atual % 60):02d}"
    index_f = str(index + 1).zfill(2)
    print(f'| Processo: {index_f} | Tempo: {duracao} | cliente: {cliente}')

print(f'===========================================================================')    
print(f'=========| Processo concluído: {index_f} carimbos gerados | Tempo:{duracao}|=====') 
print(f'===========================================================================') 

