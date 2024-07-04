from datetime import datetime
import calendar
from dateutil.relativedelta import relativedelta
from num2words import num2words
import pandas as pd
from docx import Document
from docx2pdf import convert
from docx.shared import Pt
import locale
import time
import os

# Definir a localidade para o Brasil
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

# Carregar dados do arquivo Excel
#pasta 
pasta= 'C:\\gerar_carimbos'
pasta_modelos = 'C:\\gerar_carimbos\\modelos'
#pasta_carimbo = 'C:\\carimbos_res\\carimbos'
pasta_destino = '\\g19mscp28\agroamigo\AMBIENTE\CONFORMIDADE\Time Conformidade\Renegociação\Carimbo_GED'

arquivo_excel = os.path.join(pasta,'lista_clientes_usual.xlsx')
arquivo_usual = os.path.join(pasta_modelos,'carimbo_usual.docx')
arquivo_investimento = os.path.join(pasta_modelos,'carimbo_res_inv.docx')
arquivo_custeio = os.path.join(pasta_modelos,'carimbo_res_cus.docx')

df = pd.read_excel(arquivo_excel)

quantidade_linhas = df['CLIENTE'].count()

#tempo inicial
tempo_inicio = time.time()
data_hora_atual1 = datetime.now()
data_format = data_hora_atual1.strftime("%Y/%m/%d")
data_format = f'{data_format[8:10]}/{data_format[5:7]}/{data_format[:4]}'
hora_format = data_hora_atual1.strftime("%H:%M:%S")
print(f'===========================================================================')    
print(f'=============| GERADOR DE CARIMBOS INICIADO | {data_format} {hora_format}|=====') 
print(f'===========================================================================\n')
print(f'=====[ Verificando informações...\n')

# Loop pelos clientes
for index, row in df.iterrows():
    time.sleep(1)
    if index == quantidade_linhas:
        print(f'Processo Finalizado....')
        break     
    # Extrair informações do DataFrame
    agencia = row['AGENCIA'] 
    agencia = agencia.title()
    uf = row['ESTADO']
    agencia_uf = f'{agencia}-{uf}'     
    cpf = int(row['CPF'])
    cpf = "{:011d}".format(cpf)
    cpf = f'{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}'
    cliente = row['CLIENTE']
    cliente = cliente.title()
    cliente = cliente.strip()
    contrato = row['CONTRATO']
    ccb = row['CCB']
    periodicidade = row['PERIODICIDADE'].upper()
    if periodicidade == 'ANUAL':
         text_periodicidade = 'anuais'
    elif periodicidade == 'SEMESTRAL':
        text_periodicidade = 'semestrais'
    elif periodicidade == 'TRIMESTRAL':
        text_periodicidade = 'trimestrais'

    n_parcelas = row['QTD PARCELAS']
    valor_extenso = num2words(n_parcelas, lang='pt_BR')
    if n_parcelas == 1:
         text_parcelas = 'parcela única'
    elif n_parcelas == 2:
        text_parcelas = f'2 (duas) parcelas {text_periodicidade}'      
    elif n_parcelas > 2:
        text_parcelas = f'{int(n_parcelas)} ({valor_extenso}) parcelas {text_periodicidade}'     
    
    tipo_reneg = row['TIPO_RENEG'].upper()
    if tipo_reneg == 'USUAL_CB' or tipo_reneg == 'USUAL_SB':
        if tipo_reneg == 'USUAL_CB':
            text_bonus = 'com'
        if tipo_reneg == 'USUAL_SB':     
            text_bonus = 'sem'
        modelo_word = Document(arquivo_usual)
    elif tipo_reneg == 'RESOLUCAO_I':
        modelo_word = Document(arquivo_investimento)
    elif tipo_reneg == 'RESOLUCAO_C':
        modelo_word = Document(arquivo_custeio)

    valor_renegociado = row['VALOR RENEGOCIADO']
    valor_renegociado_f = locale.format_string('%.2f', valor_renegociado, grouping=True) 
    data_renegociacao = row['DATA RENEGOCIACAO']    
    if type(data_renegociacao).__name__ == 'datetime' or type(data_renegociacao).__name__ == 'Timestamp':
        data_renegociacao = data_renegociacao.strftime('%d/%m/%Y')     
    data_parcela1 = row['PARCELA 1']
    if type(data_parcela1).__name__ == 'datetime' or type(data_parcela1).__name__ == 'Timestamp':
        data_parcela1 = data_parcela1.strftime('%d/%m/%Y')
    local_data = f'{agencia_uf}, {data_renegociacao}'         
    
    #criar a pasta da unidade
    if index == quantidade_linhas:
        print(f'Processo Finalizado....')
        break

    #verificar pastas    
    if os.path.isdir(pasta_destino):
        nome_pasta_mes = f'{pasta_destino}\{nome_pasta_mes}'     
    else:
            pasta_destino = 'C:\\gerar_carimbos\\carimbos'

    nome_pasta_mes = calendar.month_name[int(data_renegociacao[3:5])]
    nome_pasta_mes = nome_pasta_mes.upper()
    nome_pasta_dia = data_renegociacao[:2]
    nome_pasta_cliente = f'{cliente.upper()}_{contrato}_{cpf}'    
     
    nome_pasta_mes = os.path.join(pasta_destino,nome_pasta_mes)
    nome_pasta_dia = os.path.join(nome_pasta_mes, nome_pasta_dia) 
    nome_pasta_cliente = os.path.join(nome_pasta_dia, nome_pasta_cliente) 

    if not os.path.isdir(nome_pasta_mes):
         os.mkdir(nome_pasta_mes)
    if not os.path.isdir(nome_pasta_dia):
         os.mkdir(nome_pasta_dia)     
    if not os.path.isdir(nome_pasta_cliente):
         os.mkdir(nome_pasta_cliente)

    # Substituir palavras-chave no modelo do Word
    for i, paragraph in enumerate(modelo_word.paragraphs):
            #variaveis fixas(de acordo excel)
            if i > 1:
                 continue
            if i == 0:
                modelo_word.paragraphs[0].text = modelo_word.paragraphs[0].text.replace('#CLIENTE', cliente)
                modelo_word.paragraphs[0].text = modelo_word.paragraphs[0].text.replace('#CPF', cpf)
                modelo_word.paragraphs[0].text = modelo_word.paragraphs[0].text.replace('#CCB', ccb)
                modelo_word.paragraphs[0].text = modelo_word.paragraphs[0].text.replace('#DATA_ADITIVO', data_renegociacao)
                modelo_word.paragraphs[0].text = modelo_word.paragraphs[0].text.replace('#VALOR_RENEGOCIADO', f'R$ {valor_renegociado_f}')
                modelo_word.paragraphs[2].text = modelo_word.paragraphs[2].text.replace('#VALOR_RENEGOCIADO', f'R$ {valor_renegociado_f}')      
                modelo_word.paragraphs[2].text = modelo_word.paragraphs[2].text.replace('#N_PARCELAS', text_parcelas)              
                modelo_word.paragraphs[4].text = modelo_word.paragraphs[4].text.replace('#VALOR_RENEGOCIADO', f'R$ {valor_renegociado_f}')
                modelo_word.paragraphs[4].text = modelo_word.paragraphs[4].text.replace('#bonus', text_bonus)

                modelo_word.paragraphs[39].text = modelo_word.paragraphs[39].text.replace('#LOCAL_DATA', local_data) 
                if n_parcelas == 1:
                    modelo_word.paragraphs[3].text = modelo_word.paragraphs[3].text.replace(' a primeira', '')            
            #valor das parcelas

            #verificar valor par ou  impar
            #resto_centavos = ((valor_renegociado)*100) % 2    

            resto = (round(valor_renegociado,2)*100) % n_parcelas
            valor_demais_parcelas = (((valor_renegociado * 100) - resto) / n_parcelas)/100
            valor_demais_parcelas = locale.format_string('%.2f', valor_demais_parcelas, grouping=True)
            valor_demais_parcelas = f'{valor_demais_parcelas};'
            valor_ultima_parcela = float((((((valor_renegociado * 100) - resto) / n_parcelas ) + resto )/100))  
            valor_ultima_parcela = locale.format_string('%.2f', valor_ultima_parcela, grouping=True)
            valor_ultima_parcela = f'{valor_ultima_parcela}.'         
            reembolso_atual = f'Em {data_parcela1} R$ {valor_demais_parcelas}'
            if i == 0:
                modelo_word.paragraphs[6].text = modelo_word.paragraphs[6].text.replace('#REEMBOLSO1', reembolso_atual)
            if n_parcelas == 1 and '#REEMBOLSO' in modelo_word.paragraphs[7].text:
                parcela = 1
                for p in range(36,6,-1):
                        if p == 6: 
                             modelo_word.paragraphs[7]._element.clear()
                             break
                        paragraph_atual = modelo_word.paragraphs[p]
                        text_paragraph = paragraph_atual.text
                        if p < 12:                   
                            paragraph_atual._element.clear()
                        elif p >= 12 and '#REEMBOLSO' in text_paragraph:                             
                             modelo_word._element.body.remove(modelo_word.paragraphs[p]._element)
            elif n_parcelas > 1:
                    parcela = 1
                    for p in range(7,36):
                        parcela = parcela + 1                        
                        try:
                            paragraph_atual = modelo_word.paragraphs[p]
                            if periodicidade == 'SEMESTRAL':
                                if parcela == 2:
                                    data_atual = datetime.strptime(data_parcela1, "%d/%m/%Y")
                                    data_atual = data_atual +  relativedelta(months=6)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                    paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)

                                elif parcela >2:
                                    data_atual = data_atual +  relativedelta(months=6)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    if parcela == n_parcelas:
                                            reembolso_atual = f'Em {data_atual_f} R$ {valor_ultima_parcela}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                            continue
                                    elif parcela > n_parcelas and i ==0:
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
                                            if parcela > 11:
                                                if n_parcelas < 11:
                                                    remover = 18
                                                elif n_parcelas > 11:
                                                    remover = 35 - (18 - (n_parcelas-11)) 
                                                for r in range(remover,35):
                                                    if '#REEMBOLSO' in modelo_word.paragraphs[remover].text:
                                                        modelo_word._element.body.remove(modelo_word.paragraphs[remover]._element)                                           
                                                
                                    reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                    paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                            elif periodicidade == 'TRIMESTRAL':
                                if parcela == 2:
                                    data_atual = datetime.strptime(data_parcela1, "%d/%m/%Y")
                                    data_atual = data_atual +  relativedelta(months=3)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                    paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)

                                elif parcela >2:
                                    data_atual = data_atual +  relativedelta(months=3)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    if parcela == n_parcelas:
                                            reembolso_atual = f'Em {data_atual_f} R$ {valor_ultima_parcela}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                            continue
                                    elif parcela < n_parcelas and i ==0:
                                            reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                    elif parcela > n_parcelas and i ==0:                                            
                                            if parcela < 11:
                                                 paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
                                            elif parcela > 11:
                                                if n_parcelas < 11:
                                                    remover = 18
                                                elif n_parcelas > 11:
                                                    remover = 35 - (18 - (n_parcelas-11)) 
                                                for r in range(remover,35):
                                                    if '#REEMBOLSO' in modelo_word.paragraphs[remover].text:
                                                        modelo_word._element.body.remove(modelo_word.paragraphs[remover]._element)        
                                      
                            elif periodicidade == 'ANUAL':                                        
                                        data_atual = f'{data_parcela1[:6]}{int(data_parcela1[6:]) + (parcela - 1)}'
                                        if parcela == n_parcelas:
                                            reembolso_atual = f'Em {data_atual} R$ {valor_ultima_parcela}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                            continue
                                        elif parcela < n_parcelas and i ==0:
                                            reembolso_atual = f'Em {data_atual} R$ {valor_demais_parcelas}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                        elif parcela > n_parcelas and i ==0:                                            
                                            if parcela < 11:
                                                 paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
                                            elif parcela > 11:
                                                if n_parcelas < 11:
                                                    remover = 18
                                                elif n_parcelas > 11:
                                                    remover = 35 - (18 - (n_parcelas-11)) 
                                                for r in range(remover,35):
                                                    if '#REEMBOLSO' in modelo_word.paragraphs[remover].text:
                                                        modelo_word._element.body.remove(modelo_word.paragraphs[remover]._element)                     
                                        
                        except:
                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
    # Modificar a fonte de todos os estilos no documento
    for style in modelo_word.styles:
        if style.type == 1:  # type 1 significa estilo de parágrafo
            font = style.font
            font.name = 'Times New Roman'  
            font.size = Pt(11)   # Definindo o tamanho da fonte para 11 pt
    # Salvar o documento personalizado com o nome do cliente
    modelo_word.save(f'{nome_pasta_cliente}\{cliente}.docx')
    time.sleep(1)
    convert(f'{nome_pasta_cliente}\{cliente}.docx', f'{nome_pasta_cliente}\{cliente}.pdf')

    restante = quantidade_linhas - index
    tempo_atual = time.time() - tempo_inicio 
    duracao = f"{int(tempo_atual // 3600):02d}:{int((tempo_atual % 3600) // 60):02d}:{int(tempo_atual % 60):02d}"
    index_f = str(index + 1).zfill(2)
    print(f'| Processo: {index_f} | Tempo: {duracao} | cliente: {cliente}')

print(f'===========================================================================')    
print(f'=========| Processo concluído: {index_f} carimbos gerados | Tempo:{duracao}|=====') 
print(f'===========================================================================') 

