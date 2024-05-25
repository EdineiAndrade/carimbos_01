from datetime import datetime
from dateutil.relativedelta import relativedelta
import pandas as pd
from docx import Document
from docx.shared import Pt
import locale
import time
import os

# Definir a localidade para o Brasil
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

# Carregar dados do arquivo Excel
#pasta 
pasta= 'C:\\carimbos_res'
pasta_modelos = 'C:\\carimbos_res\\modelos'
pasta_carimbo = 'C:\\carimbos_res\\carimbos'

arquivo_excel = os.path.join(pasta,'lista_clientes.xlsx')
arquivo_investimento = os.path.join(pasta_modelos,'carimbo_res_inv.docx')
arquivo_custeio = os.path.join(pasta_modelos,'carimbo_res_inv.docx')

df = pd.read_excel(arquivo_excel)

# Carregar modelo do Word
modelo_investimento = Document(arquivo_investimento)
modelo_custeio = Document(arquivo_custeio)

quantidade_linhas = df.shape[0]
#tempo inicial
tempo_inicio = time.time()

#Cria uma nova pasta
##os.mkdir(pasta_reembolso)

# Loop pelos clientes
for index, row in df.iterrows():
    time.sleep(1)
    # Extrair informações do DataFrame
    agencia = row['AGENCIA'] 
    agencia = agencia.title()
    uf = row['ESTADO']
    agencia_uf = f'{agencia}-{uf}'     
    cpf = row['CPF']
    cpf = "{:011d}".format(cpf)
    cpf = f'{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}'
    cliente = row['CLIENTE']
    cliente = cliente.title()
    contrato = row['CONTRATO']
    ccb = row['CCB']
    n_parcelas = row['QTD PARCELAS']
    if n_parcelas == 1:
         text_parcelas = f'{n_parcelas} parcela'
    elif n_parcelas > 1:
        text_parcelas = f'{n_parcelas} parcelas'     
    periodicidade = row['PERIODICIDADE']
    if periodicidade == 'ANUAL':
         text_periodicidade = 'anuais'
    elif periodicidade == 'SEMESTRAL':
        text_periodicidade = 'semestrais'
    tipo_reneg = row['TIPO_RENEG']
  
    valor_renegociado = row['VALOR RENEGOCIADO']
    valor_renegociado_f = locale.format_string('%.2f', valor_renegociado, grouping=True)
    data_aditivo = row['DATA ADITIVO']
    if type(data_aditivo).__name__ == 'datetime' or type(data_aditivo).__name__ == 'Timestamp':
           data_aditivo = data_aditivo.strftime('%d/%m/%Y') 
    data_renegociacao = row['DATA RENEGOCIACAO']    
    if type(data_renegociacao).__name__ == 'datetime' or type(data_renegociacao).__name__ == 'Timestamp':
        data_renegociacao = data_renegociacao.strftime('%d/%m/%Y')     
    data_parcela1 = row['PARCELA 1']
    if type(data_parcela1).__name__ == 'datetime' or type(data_parcela1).__name__ == 'Timestamp':
        data_parcela1 = data_parcela1.strftime('%d/%m/%Y')
    local_data = f'{agencia_uf}, {data_aditivo}'         
     #data_parcela1 = f'{data_parcela1.day}/{data_parcela1.month}/{data_parcela1.year}'
    
    #criar a pasta da unidade
    if index == quantidade_linhas:
        print(f'Processo Finalizado....')
        break
   
    # Substituir palavras-chave no modelo do Word
    for i, paragraph in enumerate(modelo_investimento.paragraphs):
        if '#CLIENTE' in paragraph.text:
            paragraph.text = paragraph.text.replace('#CLIENTE', cliente)
        if '#CPF' in paragraph.text:
            paragraph.text = paragraph.text.replace('#CPF', cpf)
        if '#CCB' in paragraph.text:
            paragraph.text = paragraph.text.replace('#CCB', ccb)
        if '#DATA_ADITIVO' in paragraph.text:
            paragraph.text = paragraph.text.replace('#DATA_ADITIVO', data_aditivo)
        if '#VALOR_RENEGOCIADO' in paragraph.text:
            paragraph.text = paragraph.text.replace('#VALOR_RENEGOCIADO', f'R$ {valor_renegociado_f}')    
        if '#N_PARCELAS' in paragraph.text:
            paragraph.text = paragraph.text.replace('#N_PARCELAS', text_parcelas)
        if '#PERIODICIDADE' in paragraph.text:
            paragraph.text = paragraph.text.replace('#PERIODICIDADE', text_periodicidade)              
        if '#PARCELA1' in paragraph.text:
            paragraph.text = paragraph.text.replace('#PARCELA1', data_parcela1) 
        if '#REEMBOLSO1' in paragraph.text:            
            #n_parcelas = 10
            resto = (valor_renegociado*100) % n_parcelas
            valor_demais_parcelas = (((valor_renegociado * 100) - resto) / n_parcelas)/100
            valor_ultima_parcela = f'{(((((valor_renegociado * 100) - resto) / n_parcelas ) + resto)/100):.2f}'  
            valor_ultima_parcela = locale.format_string('%.2f', valor_demais_parcelas, grouping=True)         
            valor_demais_parcelas = locale.format_string('%.2f', valor_demais_parcelas, grouping=True)
            reembolso_atual = f'Em {data_parcela1} R$ {valor_demais_parcelas}'
            paragraph.text = paragraph.text.replace('#REEMBOLSO1', reembolso_atual)
            if n_parcelas == 1:
                parcela = 1
                for p in range(6,16):
                        parcela = parcela + 1                      
                        paragraph_atual = modelo_investimento.paragraphs[p]
                        paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
            elif n_parcelas > 1:
                    parcela = 1
                    for p in range(6,16):
                        parcela = parcela + 1                        
                        try:
                            paragraph_atual = modelo_investimento.paragraphs[p]
                            if periodicidade == 'SEMESTRAL':
                                if parcela == 2:
                                    data_atual = datetime.strptime(data_parcela1, "%d/%m/%Y")
                                    data_atual = data_atual +  relativedelta(months=6)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                    paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)

                                elif parcela >2:
                                    data_atual = data_atual +  relativedelta(months=6)
                                    data_atual_f = data_atual.strftime("%d/%m/%Y")
                                    if parcela == n_parcelas:
                                            reembolso_atual = f'Em {data_atual_f} R$ {valor_ultima_parcela}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                            continue
                                    elif parcela > n_parcelas:
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
                                            continue
                                    
                                    reembolso_atual = f'Em {data_atual_f} R$ {valor_demais_parcelas}'
                                    paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                    
                                      
                            elif periodicidade == 'ANUAL':                                        
                                        data_atual = f'{data_parcela1[:6]}{int(data_parcela1[6:]) + (parcela - 1)}'
                                        if parcela == n_parcelas:
                                            reembolso_atual = f'Em {data_atual} R$ {valor_ultima_parcela}'
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                                            continue
                                        elif parcela > n_parcelas:
                                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
                                            continue                       
                                        reembolso_atual = f'Em {data_atual} R$ {valor_demais_parcelas}'
                                        paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', reembolso_atual)
                        except:
                            paragraph_atual.text = paragraph_atual.text.replace(f'#REEMBOLSO{parcela}', '')
        if '#LOCAL_DATA' in paragraph.text:
            paragraph.text = paragraph.text.replace('#LOCAL_DATA', local_data)
        #CRONOGRAMA DE REEMBOLSO
    # Modificar a fonte de todos os estilos no documento
    for style in modelo_investimento.styles:
        if style.type == 1:  # type 1 significa estilo de parágrafo
            font = style.font
            font.name = 'Times New Roman'  # Definindo a fonte para Arial
            font.size = Pt(11)   # Definindo o tamanho da fonte para 12 pt

    # Salvar o documento personalizado com o nome do cliente
    modelo_investimento.save(f'{pasta_carimbo}\{cliente}.docx')
    modelo_investimento = Document(arquivo_investimento)
    modelo_custeio = Document(arquivo_custeio)
    restante = quantidade_linhas - index
    tempo_atual = time.time() - tempo_inicio 
    duracao = f"{int(tempo_atual // 3600):02d}:{int((tempo_atual % 3600) // 60):02d}:{int(tempo_atual % 60):02d}"
    index_f = str(index + 1).zfill(2)
    print(f'| Processo: {index_f} | Tempo: {duracao} | cliente: {cliente}')

print(f'===========================================================================')    
print(f'===| Processo concluído: {index_f} carimbos gerados | Tempo:{duracao}|=====') 
print(f'===========================================================================') 

